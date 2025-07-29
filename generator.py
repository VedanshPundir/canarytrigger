# generator.py
from docx import Document
import uuid
import requests
import os

def generate_canary_word_doc(server_url, creator_email, output_file="document.docx"):
    token_id = str(uuid.uuid4())
    
    # Register token
    try:
        response = requests.post(
            f"{server_url}/register",
            json={"token_id": token_id, "creator": creator_email}
        )
        if not response.ok:
            print("Failed to register token")
            return None
    except Exception as e:
        print(f"Registration error: {e}")
        return None
    
    # Create document
    doc = Document()
    doc.add_heading('Confidential Document', 0)
    doc.add_paragraph('This document contains sensitive information.')
    
    # Add tracking (hidden in field codes)
    tracking_url = f"{server_url}/trigger/{token_id}"
    doc.add_paragraph().add_run().add_field('INCLUDEPICTURE', f'"{tracking_url}"')
    
    # Save document
    doc.save(output_file)
    print(f"Generated canary document: {output_file}")
    return output_file

if __name__ == '__main__':
    SERVER_URL = "http://localhost:5000"  # Change to your server address
    EMAIL = "vedanshpundir43@gmail.com"      # Recipient for alerts
    generate_canary_word_doc(SERVER_URL, EMAIL)
