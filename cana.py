import threading
import tkinter as tk
from tkinter import messagebox, filedialog, scrolledtext
import requests
import webbrowser
import os
from flask import Flask, request, send_file
import datetime
import uuid
from docx import Document
import smtplib
from email.message import EmailMessage

# Flask app
app = Flask(__name__)
LOG_FILE = "alerts.log"

EMAIL_ADDRESS = "pundirved09@gmail.com"
EMAIL_PASSWORD = "pnmaayexiejdikhr"
TO_EMAIL = "pundirved09@gmail.com"

# ======================== Flask Utility =========================

def send_email_alert(subject, body):
    try:
        msg = EmailMessage()
        msg['Subject'] = subject
        msg['From'] = EMAIL_ADDRESS
        msg['To'] = TO_EMAIL
        msg.set_content(body)

        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
            smtp.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
            smtp.send_message(msg)
    except Exception as e:
        print(f"Email sending failed: {e}")

def get_location(ip):
    try:
        response = requests.get(f"https://ipinfo.io/{ip}/json", timeout=5)
        if response.status_code == 200:
            data = response.json()
            city = data.get("city", "Unknown")
            region = data.get("region", "Unknown")
            country = data.get("country", "Unknown")
            loc = data.get("loc", "Unknown")
            return f"{city}, {region}, {country} (Coordinates: {loc})"
    except Exception as e:
        return f"Location fetch error: {str(e)}"
    return "Unknown Location"

@app.route("/")
def home():
    return "<h2>✅ Canary Token Server is Running</h2>"

@app.route("/trigger/<token>")
def trigger(token):
    ip = request.headers.get("X-Forwarded-For", request.remote_addr)
    user_agent = request.headers.get('User-Agent')
    location = get_location(ip)
    timestamp = datetime.datetime.now().isoformat()

    log_line = (
        f"[{timestamp}] ALERT: Token {token} triggered\n"
        f"   - IP: {ip}\n"
        f"   - Location: {location}\n"
        f"   - User-Agent: {user_agent}\n\n"
    )
    print(log_line)

    with open(LOG_FILE, "a") as f:
        f.write(log_line)

    send_email_alert("🚨 Canary Token Triggered", log_line)
    return "📌 Token triggered and email sent."

@app.route("/generate")
def generate_url_token():
    token = str(uuid.uuid4())
    url = request.host_url + "trigger/" + token
    return f"<h3>Your Canary URL Token:</h3><br><a href='{url}'>{url}</a>"

@app.route("/word-token")
def generate_doc_token():
    token = str(uuid.uuid4())
    url = request.host_url + "trigger/" + token

    document = Document()
    document.add_heading('Company Confidential', 0)
    document.add_paragraph('Do not distribute this document. This file contains sensitive information.')
    document.add_paragraph('Monitoring Link: ' + url)

    os.makedirs("tokens", exist_ok=True)
    file_path = f"tokens/{token}.docx"
    document.save(file_path)
    return send_file(file_path, as_attachment=True)

# ======================== GUI Section =========================

def start_flask_server():
    threading.Thread(target=lambda: app.run(host="0.0.0.0", port=5000, debug=False), daemon=True).start()
    messagebox.showinfo("Server", "✅ Flask server started on http://localhost:5000")

def generate_url_token_gui():
    try:
        response = requests.get("http://localhost:5000/generate")
        webbrowser.open_new_tab("http://localhost:5000/generate")
    except Exception as e:
        messagebox.showerror("Error", str(e))

def generate_word_token_gui():
    try:
        response = requests.get("http://localhost:5000/word-token")
        with open("token.docx", "wb") as f:
            f.write(response.content)
        messagebox.showinfo("Success", "📄 Word document downloaded as token.docx")
    except Exception as e:
        messagebox.showerror("Error", str(e))

def view_logs():
    if not os.path.exists(LOG_FILE):
        messagebox.showinfo("Logs", "No logs found.")
        return
    with open(LOG_FILE, "r") as f:
        log_data = f.read()
    log_window = tk.Toplevel(root)
    log_window.title("📜 Trigger Logs")
    text = scrolledtext.ScrolledText(log_window, width=100, height=30)
    text.pack(padx=10, pady=10)
    text.insert(tk.END, log_data)

# GUI Setup
root = tk.Tk()
root.title("🛡️ Canary Token Server GUI")
root.geometry("400x300")

tk.Label(root, text="Canary Token Toolkit", font=("Arial", 16)).pack(pady=10)

tk.Button(root, text="🚀 Start Server", width=25, command=start_flask_server).pack(pady=5)
tk.Button(root, text="🔗 Generate URL Token", width=25, command=generate_url_token_gui).pack(pady=5)
tk.Button(root, text="📄 Generate Word Token", width=25, command=generate_word_token_gui).pack(pady=5)
tk.Button(root, text="📜 View Trigger Logs", width=25, command=view_logs).pack(pady=5)
tk.Button(root, text="❌ Quit", width=25, command=root.quit).pack(pady=20)

root.mainloop()
