from flask import Flask, request, send_file, render_template, flash, redirect, url_for
import os
import datetime
import uuid
from docx import Document
import requests
import smtplib
import qrcode
from email.message import EmailMessage
import sqlite3
import json

app = Flask(__name__)
app.secret_key = 'your_secret_key_here'
LOG_FILE = "alerts.log"

EMAIL_ADDRESS = ""
EMAIL_PASSWORD = ""
TO_EMAIL = ""

if not os.path.exists(LOG_FILE):
    open(LOG_FILE, 'w').close()

def send_email_alert(subject, body):
    try:
        msg = EmailMessage()
        msg['Subject'] = subject
        msg['From'] = EMAIL_ADDRESS
        msg['To'] = TO_EMAIL
        msg.set_content(body)

        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
            smtp.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
            smtp.send_message(msg)
        print("✅ Email sent")
    except Exception as e:
        print(f"❌ Email error: {e}")

def get_location(ip):
    try:
        response = requests.get(f"https://ipinfo.io/{ip}/json", timeout=5)
        if response.status_code == 200:
            data = response.json()
            loc = data.get("loc", "0,0")
            city = data.get("city", "Unknown")
            region = data.get("region", "Unknown")
            country = data.get("country", "Unknown")
            return {
                "text": f"{city}, {region}, {country}",
                "loc": loc
            }
        return {"text": "Unknown", "loc": "0,0"}
    except Exception as e:
        print(f"Location lookup failed: {e}")
        return {"text": "Unknown", "loc": "0,0"}

def parse_alerts():
    alerts = []
    with open(LOG_FILE, "r") as f:
        blocks = f.read().strip().split("\n\n")
        for entry in blocks:
            lines = entry.strip().split("\n")
            try:
                timestamp = lines[0].split("]")[0][1:]
                token = lines[0].split("Token ")[1]
                ip = lines[1].split("IP: ")[1]
                location_line = lines[2].split("Location: ")[1]
                user_agent = lines[3].split("User-Agent: ")[1]
                
                if "(Coordinates:" in location_line:
                    location_text = location_line.split(" (Coordinates:")[0]
                    coords = location_line.split("Coordinates: ")[1].strip(")")
                    lat, lon = coords.split(",")
                else:
                    location_text = location_line
                    lat, lon = "0", "0"

                alerts.append({
                    "timestamp": timestamp,
                    "token": token,
                    "ip": ip,
                    "location": location_text,
                    "user_agent": user_agent,
                    "latitude": lat,
                    "longitude": lon
                })
            except Exception as e:
                print(f"Parse error: {e}")
                continue
    return alerts
SPLUNK_HEC_URL = "http://localhost:8088/services/collector/event"
SPLUNK_TOKEN = "********-b21e-48e2-*****-************"

def send_to_splunk(token, ip, location, user_agent, message, timestamp, username=None, password=None, success=None, attempts=None):
    payload = {
        "event": {
            "token": token,
            "ip": ip,
            "location": location,
            "user_agent": user_agent,
            "message": message,
            "timestamp": timestamp,
            "username": username,
            "password": password,
            "login_success": success,
            "login_attempts": attempts
        },
        "sourcetype": "_json"
    }

    headers = {
        "Authorization": f"Splunk {SPLUNK_TOKEN}"
    }

    try:
        res = requests.post(SPLUNK_HEC_URL, headers=headers, data=json.dumps(payload), verify=False)
        if res.status_code != 200:
            print("Splunk HEC Error:", res.text)
    except Exception as e:
        print("Error sending to Splunk:", e)

@app.route('/alerts-count')
def alerts_count():
    count = 0
    try:
        with open("alerts.log", "r") as file:
            content = file.read().strip()
            if content:
                count = len(content.split("\n\n"))
    except Exception as e:
        print(f"Error reading alerts.log: {e}")
    return {"count": count}

@app.route("/")
def home():
    alerts = parse_alerts()
    return render_template("home.html",
                           url_tokens=len([a for a in alerts if a["token"].startswith("http")]),
                           doc_tokens=len([a for a in alerts if not a["token"].startswith("http")]),
                           alerts=len(alerts))

@app.route("/generate")
def generate_url_token():
    token = str(uuid.uuid4())
    url = request.host_url + "trigger/" + token
    return render_template("generate_token.html", url=url)

@app.route("/word-token")
def generate_doc_token():
    token = str(uuid.uuid4())
    url = request.host_url + "trigger/" + token

    # Create Word Document
    document = Document()
    document.add_heading('Company Confidential', 0)
    document.add_paragraph('This document is sensitive and should only be accessed by the legitimate user.')

    # Generate QR Code
    qr = qrcode.make(url)
    os.makedirs("qrcodes", exist_ok=True)
    qr_path = f"qrcodes/{token}.png"
    qr.save(qr_path)

    # Insert QR Code into Word Document
    document.add_paragraph("Scan this QR code to access:")
    document.add_picture(qr_path)

    # Save DOCX
    os.makedirs("tokens", exist_ok=True)
    file_path = f"tokens/{token}.docx"
    document.save(file_path)

    return send_file(file_path, as_attachment=True)

    return send_file(file_path, as_attachment=True)

@app.route("/trigger/<token>", methods=["GET", "POST"])
def trigger(token):
    ip = request.headers.get("X-Forwarded-For", request.remote_addr)
    user_agent = request.headers.get("User-Agent", "Unknown")
    location_info = get_location(ip)
    timestamp = datetime.datetime.now().isoformat()
    lat, lon = location_info['loc'].split(",")

    conn = sqlite3.connect("alerts.db")
    c = conn.cursor()

    # Create vulnerable login table
    c.execute('''
        CREATE TABLE IF NOT EXISTS login_honeypot (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            token TEXT,
            ip TEXT,
            location TEXT,
            latitude TEXT,
            longitude TEXT,
            user_agent TEXT,
            username TEXT,
            password TEXT,
            success INTEGER,
            timestamp TEXT
        )
    ''')

    # Count attempts from this IP/token
    c.execute("SELECT COUNT(*) FROM login_honeypot WHERE ip = ? AND token = ?", (ip, token))
    attempt_count = c.fetchone()[0]

    success = False
    username = password = ""

    if request.method == "POST":
        username = request.form.get("username", "")
        password = request.form.get("password", "")

        # INTENTIONAL VULNERABLE QUERY
        query = f"SELECT * FROM users WHERE username = '{username}' AND password = '{password}'"

        try:
            c.execute('CREATE TABLE IF NOT EXISTS users (username TEXT, password TEXT)')
            c.execute("INSERT INTO users (username, password) VALUES ('admin', 'admin123')")  # dummy
            conn.commit()
            c.execute(query)
            result = c.fetchone()
            if result:
                success = True
        except Exception as e:
            print("SQL Error:", e)

        # Log to DB
        c.execute('''
            INSERT INTO login_honeypot (token, ip, location, latitude, longitude, user_agent, username, password, success, timestamp)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (token, ip, location_info["text"], lat, lon, user_agent, username, password, int(success), timestamp))
        conn.commit()

        # Log to alerts.log
        log_line = (
            f"[{timestamp}] Token {token}\n"
            f"IP: {ip}\n"
            f"Location: {location_info['text']} (Coordinates: {lat},{lon})\n"
            f"User-Agent: {user_agent}\n"
            f"Username: {username} | Password: {password}\n"
            f"Login Success: {success}\n\n"
        )

        with open(LOG_FILE, "a", encoding="utf-8") as f:
            f.write(log_line)

        # Send alerts
        send_email_alert("🚨 Login Honeypot Triggered", log_line)
        send_to_splunk(
            token=token,
            ip=ip,
            location=location_info["text"],
            user_agent=user_agent,
            message=f"{username}:{password}",
            timestamp=timestamp,
            username=username,
            password=password,
            success=success,
            attempts=attempt_count + 1 if not success else attempt_count
        )

    conn.close()

    return render_template("login.html", success=success, attempts=(attempt_count + 1 if not success else attempt_count))

    conn.close()

    return render_template("login.html", success=success, attempts=(attempt_count + 1 if not success else attempt_count))
"""@app.route("/alerts")
def view_alerts():
    alerts = parse_alerts()

    # Load confessions
    conn = sqlite3.connect("alerts.db")
    c = conn.cursor()
    c.execute("SELECT token, ip, location, latitude, longitude, user_agent, message, timestamp FROM login_honeypot ORDER BY id DESC")
    confessions = c.fetchall()
    conn.close()



    return render_template("alerts.html", alerts=alerts[::-1], confessions=confessions)"""

@app.route("/alerts")
def view_alerts():
    alerts = parse_alerts()

    # Load and group confessions
    conn = sqlite3.connect("alerts.db")
    c = conn.cursor()
    c.execute("SELECT token, message, timestamp FROM confessions ORDER BY id DESC")
    confessions = c.fetchall()
    conn.close()

    grouped_confessions = {}
    for token, message, timestamp in confessions:
        if token not in grouped_confessions:
            grouped_confessions[token] = []
        grouped_confessions[token].append({
            "message": message,
            "timestamp": timestamp
        })

    return render_template("alerts.html", alerts=alerts[::-1], grouped_confessions=grouped_confessions)

@app.route("/clear-alerts")
def clear_alerts():
    open(LOG_FILE, 'w').close()
    flash("Alerts cleared.", "success")
    return redirect(url_for("view_alerts"))

@app.route("/map")
def show_map():
    return render_template("map.html")

if __name__ == "__main__":
    app.run(debug=True)
